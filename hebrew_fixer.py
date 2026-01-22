"""
Hebrew Text Fixer for Affinity Canvas
Reverses Hebrew words individually while preserving word order and handling nikud marks correctly.
"""
import sys
import os
import unicodedata
import regex
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout,
    QPushButton, QTextEdit, QCheckBox, QComboBox,
    QLabel, QFileDialog, QMessageBox
)
from PyQt5.QtGui import QFont
from PyQt5.QtCore import Qt

try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from odf.opendocument import OpenDocumentText, load
    from odf.text import P, Span
    ODT_AVAILABLE = True
except ImportError:
    ODT_AVAILABLE = False


def reverse_hebrew_word(word: str) -> str:
    """
    Reverse a single Hebrew word, handling nikud (diacritic marks) correctly.
    
    Strategy:
    1. Split word into base character + combining marks units
    2. Reverse the order of units
    3. Within each unit, place marks before the base character
    
    Args:
        word: A single Hebrew word (no spaces)
    
    Returns:
        The reversed word with nikud marks properly positioned
    """
    # Normalize to NFC form
    word = unicodedata.normalize("NFC", word)
    
    # Match Hebrew base character followed by any combining marks
    units = regex.findall(r'\p{Hebrew}\p{M}*', word, flags=regex.UNICODE)
    
    if not units:
        return word
    
    # Reverse the order of units and flip marks+base in each unit
    flipped = []
    for unit in units:
        base = unit[0]
        marks = unit[1:]
        # Place marks before base character for proper rendering
        flipped.append(marks + base)
    
    return "".join(flipped[::-1])


def process_text(text: str, do_reverse: bool = True) -> str:
    """
    Process text by reversing Hebrew words AND word order for Affinity Canvas.
    
    Affinity Canvas reverses Hebrew text, so we need to:
    1. Reverse each individual word (with nikud handling)
    2. Reverse the order of words in each Hebrew phrase
    
    Args:
        text: Input text containing Hebrew and possibly other languages
        do_reverse: Whether to reverse Hebrew words (for preview toggle)
    
    Returns:
        Processed text with Hebrew words and word order reversed
    """
    if not do_reverse:
        return text
    
    # Pattern to match entire Hebrew phrases (multiple words together)
    # Split on non-Hebrew boundaries to get Hebrew phrases
    hebrew_phrase_pattern = regex.compile(
        r'((?:\p{Hebrew}[\p{Hebrew}\p{M}\s]*)+)',
        flags=regex.UNICODE
    )
    
    def process_phrase(match):
        phrase = match.group(1)
        
        # Extract individual Hebrew words from the phrase
        hebrew_word_pattern = regex.compile(r'(\p{Hebrew}[\p{Hebrew}\p{M}]*)', flags=regex.UNICODE)
        words = hebrew_word_pattern.findall(phrase)
        
        if not words:
            return phrase
        
        # Reverse each word individually
        reversed_words = [reverse_hebrew_word(word) for word in words]
        
        # Reverse the order of words
        reversed_words = reversed_words[::-1]
        
        # Get the spacing/punctuation between words to preserve it
        spaces = regex.split(r'\p{Hebrew}[\p{Hebrew}\p{M}]*', phrase, flags=regex.UNICODE)
        spaces = [s for s in spaces if s]  # Remove empty strings
        
        # Reconstruct with reversed word order
        result = []
        for i, word in enumerate(reversed_words):
            result.append(word)
            if i < len(spaces):
                result.append(spaces[min(i, len(spaces)-1)])
        
        return ''.join(result)
    
    return hebrew_phrase_pattern.sub(process_phrase, text)


def load_plain_text(path: str) -> str:
    """Load text from various file formats."""
    ext = os.path.splitext(path)[1].lower()
    
    if ext == ".txt":
        with open(path, encoding="utf-8") as f:
            return f.read()
    
    elif ext == ".docx":
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)
    
    elif ext == ".odt":
        if not ODT_AVAILABLE:
            raise ImportError("odfpy not installed. Install with: pip install odfpy")
        odt = load(path)
        lines = []
        for p in odt.getElementsByType(P):
            txt = ""
            for c in p.childNodes:
                if c.nodeType == c.TEXT_NODE:
                    txt += c.data
                elif isinstance(c, Span) and c.firstChild:
                    txt += c.firstChild.data
            lines.append(txt)
        return "\n".join(lines)
    
    elif ext == ".rtf":
        # RTF handling would require additional library (striprtf or similar)
        raise ValueError("RTF format requires additional library. Save as TXT or DOCX instead.")
    
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def process_docx(in_path: str, out_path: str, do_reverse: bool):
    """Process DOCX file while preserving formatting."""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not installed")
    
    src = Document(in_path)
    dst = Document()
    
    for para in src.paragraphs:
        newp = dst.add_paragraph()
        newp.alignment = para.alignment
        newp.paragraph_format.right_to_left = True  # Set RTL for Hebrew
        
        for run in para.runs:
            text = process_text(run.text, do_reverse)
            nr = newp.add_run(text)
            nr.bold = run.bold
            nr.italic = run.italic
            nr.underline = run.underline
            nr.font.name = run.font.name or "Times New Roman"
            nr.font.size = run.font.size or Pt(12)
    
    dst.save(out_path)


def process_odt(in_path: str, out_path: str, do_reverse: bool):
    """Process ODT file while preserving formatting."""
    if not ODT_AVAILABLE:
        raise ImportError("odfpy not installed")
    
    src = load(in_path)
    dst = OpenDocumentText()
    
    # Copy styles and auto-styles
    for style in src.styles.childNodes:
        dst.styles.addElement(style)
    for astyle in src.automaticstyles.childNodes:
        dst.automaticstyles.addElement(astyle)
    
    for p in src.getElementsByType(P):
        sty = p.getAttribute("stylename")
        newp = P(stylename=sty)
        
        for c in p.childNodes:
            if c.nodeType == c.TEXT_NODE:
                newp.addText(process_text(c.data, do_reverse))
            elif isinstance(c, Span):
                t = c.firstChild.data if c.firstChild else ""
                sname = c.getAttribute("stylename")
                sp = Span(stylename=sname, text=process_text(t, do_reverse))
                newp.addElement(sp)
        
        dst.text.addElement(newp)
    
    dst.save(out_path)


class HebrewFixerWindow(QMainWindow):
    """Main application window."""
    
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Hebrew Text Fixer for Affinity Canvas")
        self.resize(900, 700)
        
        self.input_path = None
        self.raw_text = ""
        self.do_reverse = True
        
        self._setup_ui()
    
    def _setup_ui(self):
        """Set up the user interface."""
        central = QWidget()
        layout = QVBoxLayout()
        
        # Title label
        title = QLabel("Hebrew Text Fixer")
        title.setFont(QFont("Arial", 16, QFont.Bold))
        layout.addWidget(title)
        
        # Instructions
        instructions = QLabel(
            "Fix Hebrew text for Affinity Canvas by reversing words AND word order.\n"
            "Both individual words and their sequence will be reversed."
        )
        instructions.setWordWrap(True)
        layout.addWidget(instructions)
        
        # Open file button
        btn_open = QPushButton("📂 Open File...")
        btn_open.clicked.connect(self.open_file)
        layout.addWidget(btn_open)
        
        # Reverse checkbox
        self.cb_reverse = QCheckBox("Reverse Hebrew words (fix for Affinity Canvas)")
        self.cb_reverse.setChecked(True)
        self.cb_reverse.stateChanged.connect(self.update_preview)
        layout.addWidget(self.cb_reverse)
        
        # Output format selector
        layout.addWidget(QLabel("Save as:"))
        self.format_cb = QComboBox()
        formats = ["txt"]
        if DOCX_AVAILABLE:
            formats.append("docx")
        if ODT_AVAILABLE:
            formats.append("odt")
        self.format_cb.addItems(formats)
        layout.addWidget(self.format_cb)
        
        # Preview button
        self.btn_preview = QPushButton("🔄 Preview")
        self.btn_preview.clicked.connect(self.update_preview)
        self.btn_preview.setEnabled(False)
        layout.addWidget(self.btn_preview)
        
        # Text editor for preview
        self.editor = QTextEdit()
        self.editor.setReadOnly(True)
        self.editor.setFont(QFont("David", 14))  # Hebrew font
        self.editor.setLayoutDirection(Qt.RightToLeft)
        self.editor.setAlignment(Qt.AlignRight)
        self.editor.setPlaceholderText("Processed text will appear here...")
        layout.addWidget(self.editor)
        
        # Save button
        self.btn_save = QPushButton("💾 Save Output...")
        self.btn_save.clicked.connect(self.save_output)
        self.btn_save.setEnabled(False)
        layout.addWidget(self.btn_save)
        
        # Status label
        self.status_label = QLabel("Ready. Open a file to begin.")
        layout.addWidget(self.status_label)
        
        central.setLayout(layout)
        self.setCentralWidget(central)
    
    def open_file(self):
        """Open and load a file."""
        path, _ = QFileDialog.getOpenFileName(
            self,
            "Select File",
            "",
            "All Supported (*.txt *.docx *.odt);;Text Files (*.txt);;Word Documents (*.docx);;OpenDocument (*.odt)"
        )
        
        if not path:
            return
        
        try:
            self.raw_text = load_plain_text(path)
            self.input_path = path
            self.status_label.setText(f"Loaded: {os.path.basename(path)}")
            self.btn_preview.setEnabled(True)
            self.update_preview()
        except Exception as e:
            QMessageBox.critical(self, "Error Loading File", str(e))
            self.status_label.setText("Error loading file")
    
    def update_preview(self):
        """Update the preview with processed text."""
        if not self.input_path:
            return
        
        self.do_reverse = self.cb_reverse.isChecked()
        processed = process_text(self.raw_text, self.do_reverse)
        self.editor.setPlainText(processed)
        self.btn_save.setEnabled(True)
        self.status_label.setText("Preview updated. Ready to save.")
    
    def save_output(self):
        """Save the processed text to a file."""
        fmt = self.format_cb.currentText()
        
        # Set up file filter
        filters = {
            "txt": "Text Files (*.txt)",
            "docx": "Word Documents (*.docx)",
            "odt": "OpenDocument (*.odt)"
        }
        
        path, _ = QFileDialog.getSaveFileName(
            self,
            "Save Processed File",
            "",
            filters.get(fmt, "All Files (*.*)")
        )
        
        if not path:
            return
        
        # Ensure correct extension
        if not path.lower().endswith(f".{fmt}"):
            path += f".{fmt}"
        
        try:
            if fmt == "txt":
                with open(path, "w", encoding="utf-8") as f:
                    f.write(process_text(self.raw_text, self.do_reverse))
            
            elif fmt == "docx":
                if not DOCX_AVAILABLE:
                    raise ImportError("python-docx not installed")
                process_docx(self.input_path, path, self.do_reverse)
            
            elif fmt == "odt":
                if not ODT_AVAILABLE:
                    raise ImportError("odfpy not installed")
                process_odt(self.input_path, path, self.do_reverse)
            
            QMessageBox.information(self, "Success", f"File saved successfully:\n{path}")
            self.status_label.setText(f"Saved: {os.path.basename(path)}")
        
        except Exception as e:
            QMessageBox.critical(self, "Error Saving File", str(e))
            self.status_label.setText("Error saving file")


def main():
    """Main entry point."""
    app = QApplication(sys.argv)
    window = HebrewFixerWindow()
    window.show()
    sys.exit(app.exec_())


if __name__ == "__main__":
    main()
